import streamlit as st
import whisper
import os
from pydub import AudioSegment
from fpdf import FPDF
from docx import Document
from io import BytesIO
import tempfile

# --- CONFIGURAÇÃO DA PÁGINA ---
st.set_page_config(
    page_title="Transcritor Inteligente",
    page_icon="🎙️",
    layout="centered"
)

# --- ESTILIZAÇÃO UX (CSS) ---
st.markdown("""
    <style>
    .main { border-radius: 10px; }
    .stButton>button { width: 100%; border-radius: 5px; height: 3em; background-color: #4CAF50; color: white; }
    .stDownloadButton>button { width: 100%; border-radius: 5px; height: 3em; }
    </style>
    """, unsafe_allow_html=True)

# --- FUNÇÕES DE BACKEND ---

@st.cache_resource
def load_model():
    """Carrega o modelo Whisper uma única vez para otimizar performance."""
    # Mudamos de 'base' para 'small' para maior precisão em português
    return whisper.load_model("small")

def process_audio_to_wav(input_file):
    """Converte qualquer formato suportado para .wav para garantir compatibilidade."""
    try:
        audio = AudioSegment.from_file(input_file)
        temp_wav = tempfile.NamedTemporaryFile(delete=False, suffix=".wav")
        audio.export(temp_wav.name, format="wav")
        return temp_wav.name
    except Exception as e:
        st.error(f"Erro ao processar áudio: {e}")
        return None

def generate_pdf(text):
    """Gera um arquivo PDF formatado."""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    # Substituir caracteres não-latin1 para evitar erros no FPDF simples
    clean_text = text.encode('latin-1', 'replace').decode('latin-1')
    pdf.multi_cell(0, 10, clean_text)
    return pdf.output()

def generate_docx(text):
    """Gera um arquivo Word formatado."""
    doc = Document()
    doc.add_heading('Transcrição de Áudio', 0)
    doc.add_paragraph(text)
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

def mock_google_docs_upload(text):
    """
    Simula a integração com a Google Docs API.
    Em um cenário real, aqui seria implementado o fluxo de credenciais OAuth2.
    """
    # Exemplo de como a estrutura seria iniciada:
    # creds = service_account.Credentials.from_json_keyfile_name(SERVICE_ACCOUNT_FILE, SCOPES)
    # service = build('docs', 'v1', credentials=creds)
    st.info("Simulando autenticação OAuth2 com Google Cloud...")
    st.success("Documento criado com sucesso no seu Google Drive (Simulação)!")

# --- INTERFACE DE USUÁRIO (UI) ---

def main():
    st.title("🎙️ Transcritor de Áudio Especialista")
    st.markdown("Converta suas gravações em texto de forma simples e profissional.")

    # Opções de Entrada
    tab1, tab2 = st.tabs(["📁 Upload de Arquivo", "🔗 Link do Áudio"])
    
    audio_source = None

    with tab1:
        uploaded_file = st.file_uploader("Arraste ou selecione seu áudio", 
                                        type=['mp3', 'wav', 'ogg', 'flac', 'm4a'])
        if uploaded_file:
            audio_source = uploaded_file

    with tab2:
        audio_url = st.text_input("Cole o link direto do arquivo de áudio (URL)")
        if audio_url:
            st.info("Nota: O link deve apontar diretamente para um arquivo de áudio público.")
            audio_source = audio_url

    if audio_source:
        if st.button("🚀 Iniciar Transcrição"):
            model = load_model()
            
            with st.spinner("⏳ Processando e convertendo áudio..."):
                # Salvando arquivo temporário para processamento
                with tempfile.NamedTemporaryFile(delete=False, suffix=".tmp") as tmp_file:
                    if isinstance(audio_source, str): # Se for URL (Simplificado para este exemplo)
                        # Aqui usaria requests para baixar, mas usaremos o arquivo local para o MVP
                        st.error("Funcionalidade de URL requer implementação de download. Use Upload.")
                        return
                    else:
                        tmp_file.write(audio_source.read())
                        audio_path = tmp_file.name

                # Padronização de formato
                wav_path = process_audio_to_wav(audio_path)   
    
                if wav_path:
                    try:
                        st.info("🧠 Inteligência Artificial transcrevendo... Forçando idioma: Português. Isso pode levar alguns minutos dependendo do tamanho.")
                        # Adicionamos parâmetros para evitar alucinações:
                        # 1. language='pt' (Garante que ele não tente adivinhar Coreano/Inglês)
                        # 2. fp16=False (Melhora estabilidade em CPUs Windows)
                        # 3. beam_size=5 (Melhora a busca pelas palavras corretas)
                        result = model.transcribe(
                            wav_path, 
                            language='pt', 
                            fp16=False,
                            beam_size=5 
                        )
                        
                        transcription_text = result['text']
                        
                        st.session_state['transcription'] = transcription_text
                        st.success("✅ Transcrição concluída!")
                        
                    except Exception as e:
                        st.error(f"Erro durante a transcrição: {e}")
                    finally:
                        if os.path.exists(wav_path): os.remove(wav_path)
                        if os.path.exists(audio_path): os.remove(audio_path)

    # Exibição dos Resultados e Exportação
    if 'transcription' in st.session_state:
        st.divider()
        st.subheader("📝 Resultado da Transcrição")
        text_area = st.text_area("Texto extraído:", st.session_state['transcription'], height=300)

        st.subheader("📥 Exportar Documento")
        col1, col2, col3 = st.columns(3)

        with col1:
            pdf_bytes = generate_pdf(text_area)
            st.download_button(
                label="📄 Baixar PDF",
                data=bytes(pdf_bytes),
                file_name="transcricao.pdf",
                mime="application/pdf"
            )

        with col2:
            docx_bytes = generate_docx(text_area)
            st.download_button(
                label="📘 Baixar DOCX",
                data=docx_bytes,
                file_name="transcricao.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        with col3:
            if st.button("☁️ Google Docs"):
                mock_google_docs_upload(text_area)

if __name__ == "__main__":
    main()